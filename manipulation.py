# coding:utf-8
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
import numpy as np
import os
import sys
# import seaborn as sns  # 学院派风格的话，这个用不到
import docx   # 注意是python-docx
from docx.shared import Inches
matplotlib.style.use('ggplot')

plt.rcParams["font.sans-serif"] = ["Microsoft Yahei"]
plt.rcParams["axes.unicode_minus"] = False

# sns.set(color_codes=True)
# sns.set(rc={'font.family': [u'Microsoft YaHei']})
# sns.set(rc={'font.sans-serif': [u'Microsoft YaHei', u'Arial',
#                                 u'Liberation Sans', u'Bitstream Vera Sans',
#                                 u'sans-serif']})


def frc(std, side, how):
    return ("L_{std}H_{side}RLFCFBK{how}"
            .format(std=std, side=side, how=how))


def pos(std, side, access):
    return ("L_{std}H_{side}{access}POSFBK"
            .format(std=std, side=side, access=access))


def slp(side, access, how, times):
    return ("k_{side}_{access}_{how}_{times}"
            .format(side=side, access=access, how=how, times=times))


def box_col(std, the_date):
    return ("{std}_{the_date}"
            .format(std=std, the_date=the_date))


def progressbar(i):
    """ 小型的进度条 """
    sys.stdout.write('第{0}机架计算中！- 总共7机架\r'.format(i + 1))
    sys.stdout.flush()


def fst_idx(s, level_line):
    return s.loc[s > level_line].sort_index().index[0]


def lst_idx(s):

    return s.sort_values().index[-1]


def cut(df, s, level_line):
    if s.loc[s > level_line].shape[0] == 0:
        return df.loc[s > level_line]
    else:
        # print(s[s.loc[s > level_line].index[0]],
        #      s[s.sort_values().index[-1]])
        return df.loc[(s.index >= fst_idx(s, level_line)) &
                      (s.index < lst_idx(s))]


# --- setup para ---
# data directory
root_dir = "../data/零调数据"

date_num_list = [-1, -2, -3]
times_id_list = ["A", "B", "C", "D", "E", "F"]
std_list = ["F1", "F2", "F3", "F4", "F5", "F6", "F7"]
side_list = ["OS", "DS"]
access_list = ["ENT", "EXT"]
how_list = ["LC", "PT"]
level_line = 4000
total_line = 3000   # 保持率计算

last_num = -1
last_date = os.listdir(root_dir)[last_num]

box_data = pd.DataFrame()

for date_num in date_num_list:
    summary = pd.DataFrame()
    the_date = os.listdir(root_dir)[date_num]
    data_dir = "/".join([root_dir, the_date])
    for times_id in times_id_list:
        # print(times_id)
        file_name = [x for x in os.listdir(data_dir) if times_id in x][0]
        df = pd.read_csv("/".join([data_dir, file_name]))
        for std in std_list:
            # selected series 1@
            # selected_s = df[frc(std, "OS", "LC")]
            # df = cut(df, selected_s, level_line)
            # df = df.loc[df[frc(std, "OS", "LC")] > level_line]
            for side in side_list:
                for access in access_list:
                    for how in how_list:
                        # selected series 2@
                        # selected_s = df[frc(std, side, how)]
                        # df = cut(df, selected_s, level_line)
                        df = df.loc[df[frc(std, side, how)] > level_line]
                        pos_size = df[pos(std, side, access)].shape[0]
                        frc_size = df[frc(std, side, how)].shape[0]
                        if ((pos_size != 0) and (frc_size != 0) and
                                (pos_size == frc_size)):
                            summary.loc[std,
                                        slp(side, access, how, times_id)] = (
                                np.polyfit(df[pos(std, side, access)],
                                           df[frc(std, side, how)], 1)[0]
                            )
                        else:
                            summary.loc[std, slp(
                                side, access, how, times_id)] = np.nan

    if not os.path.exists("../data/inter/{}".format(the_date)):
        os.makedirs("../data/inter/{}".format(the_date))
    if not os.path.exists("../data/pic/{}".format(the_date)):
        os.makedirs("../data/pic/{}".format(the_date))
    summary.to_excel("../data/inter/{}/中间计算结果.xlsx".format(the_date))

    plot_data = pd.DataFrame()
    for how in how_list:
        for std in std_list:
            the_col_list = [x for x in summary.columns if how in x]
            the_series = summary.loc[std, the_col_list]
            the_series = the_series.loc[the_series > 1000]
            if the_series.shape[0] != 0:
                plot_data.loc[std, how] = round(
                    the_series.mean() / total_line * 100, 2)
            else:
                plot_data.loc[std, how] = np.nan

    # --- plt.fig ---
    plt.figure(0)
    plot_data.plot(kind="bar")
    plt.title = "LC和PT刚度比较"
    plt.xlabel("机架")
    plt.ylabel("刚度保持率%")

    x = np.arange(len(plot_data.index))
    y1 = plot_data["LC"]
    y2 = plot_data["PT"]
    for a, b in zip(x, y1):
        plt.text(a - 0.3, b + 0.05, '%.1f' %
                 b, ha='center', va='bottom', fontsize=10)
    for a, b in zip(x, y2):
        plt.text(a + 0.3, b + 0.05, '%.1f' %
                 b, ha='center', va='bottom', fontsize=10)
    plt.savefig("../data/pic/{}/bar_plot.png".format(the_date))
    plt.close(0)

    # 全机架保持率计算
    how = "LC"
    for times_id in times_id_list:
        for std in std_list:
            the_col_list = [
                x for x in summary.columns if times_id in x if how in x]
            the_series = summary.loc[std, the_col_list]
            the_series = the_series.loc[the_series > 0.001]
            if the_series.shape[0] != 0:
                box_data.loc[times_id, box_col(std, the_date)] = round(
                    the_series.mean() / total_line * 100, 2)
            else:
                box_data.loc[times_id, box_col(std, the_date)] = np.nan
                # or 0

box_data.to_excel("../data/inter/{}/box计算结果.xlsx".format(last_date))

for std in std_list:
    the_col_list = [
        x for x in box_data.columns if std in x]
    plt.figure(0)
    box_data[the_col_list].boxplot()
    # print(std)
    plt.title = "{}机架刚度保持率".format(std)
    plt.xlabel("日期时间")
    plt.ylabel("刚度保持率%")
    plt.savefig("../data/pic/{}/box_plot_{}.png".format(last_date, std))
    plt.close(0)

doc = docx.Document("base.docx")
doc.add_paragraph("刚度评价{}".format(last_date))
doc.paragraphs[-1].style = "Title"

doc.add_paragraph("    LC和PT刚度如下图所示。")
doc.add_picture("../data/pic/{}/bar_plot.png".format(last_date),
                width=Inches(6.2))
for std in std_list:
    doc.add_paragraph("   {}机架近3周刚度保持率如下图所示。".format(std))
    doc.add_picture(
        "../data/pic/{}/box_plot_{}.png".format(last_date, std),
        width=Inches(6.2))

if not os.path.exists("../data/result"):
    os.makedirs("../data/result")
doc.save("../data/result/轧机刚度评价{}.docx".format(last_date))
